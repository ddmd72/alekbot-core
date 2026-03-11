"""
RichContentService — application-layer dispatcher for rich content delivery.

Agents declare *what* to generate (RichContent); this service handles *how*:
  - converts LLM-generated text to binary formats (xlsx, docx)
  - .html files → GCS public URL (returned to caller for link delivery)
  - other files → direct platform upload via PlatformMediaPort

ConversationHandler calls process() for each non-table rich content item
after the text response has been delivered.

Returns:
  Optional[str] — public URL when content was stored in GCS (html),
                  None when content was uploaded as a file attachment.

Supported types:
  file      — LLM-generated content
              .html              → GCS upload → return URL
              .md / .txt         → UTF-8 encode → platform upload
              .xlsx (CSV→xlsx)   → platform upload
              .docx (Markdown→docx) → platform upload
  widget    — Agent-generated HTML → Playwright screenshot → PNG → platform upload (inline image)

Planned (M3+):
  map_image — Google Maps Static API → GCS upload → return URL
"""
import csv
import io
import uuid
from typing import Optional

from ..ports.platform_media_port import PlatformMediaPort
from ..ports.media_storage_port import MediaStoragePort
from ..ports.html_renderer_port import HtmlRendererPort, HtmlRenderError
from ..domain.messaging import RichContent
from ..utils.logger import logger


class RichContentService:
    """Converts and delivers rich content items via PlatformMediaPort / MediaStoragePort."""

    def __init__(
        self,
        media_port: PlatformMediaPort,
        storage_port: Optional[MediaStoragePort] = None,
        html_renderer: Optional[HtmlRendererPort] = None,
    ) -> None:
        self._media_port = media_port
        self._storage_port = storage_port
        self._html_renderer = html_renderer

    async def process(self, content: RichContent, channel_id: str) -> Optional[str]:
        """
        Process a single RichContent item and deliver it to the platform.

        Args:
            content:    Structured content descriptor from LLM output
            channel_id: Platform-specific channel identifier

        Returns:
            Public URL string if content was stored in GCS (html file type),
            None for all other types (direct upload or skip).
        """
        if content.content_type == "file":
            return await self._handle_file(content, channel_id)
        elif content.content_type == "widget":
            await self._handle_widget(content, channel_id)
            return None
        else:
            logger.warning(
                "RichContentService: unsupported content type '%s' — skipping",
                content.content_type,
            )
            return None

    # ------------------------------------------------------------------
    # Private handlers
    # ------------------------------------------------------------------

    async def _handle_file(self, content: RichContent, channel_id: str) -> Optional[str]:
        filename = content.data.get("filename", "document.md").strip()
        title = content.data.get("title", filename).strip()
        text_content = content.data.get("content", "")

        if not text_content:
            logger.warning("RichContentService: file type missing 'content' in data")
            return None

        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "md"

        # HTML → GCS public URL (if storage_port configured)
        if ext == "html" and self._storage_port:
            return await self._store_html(text_content, filename)

        # All other formats → encode/convert → platform file upload
        try:
            if ext == "xlsx":
                file_bytes = _csv_to_xlsx(text_content)
            elif ext == "docx":
                file_bytes = _markdown_to_docx(text_content)
            else:
                file_bytes = text_content.encode("utf-8")
        except Exception as e:
            logger.error(
                "RichContentService: conversion failed for '%s' — %s; falling back to plain text",
                filename,
                e,
            )
            file_bytes = text_content.encode("utf-8")
            filename = filename.rsplit(".", 1)[0] + ".txt"

        await self._media_port.upload_file(
            file_bytes=file_bytes,
            filename=filename,
            title=title,
            channel_id=channel_id,
        )
        return None

    async def _handle_widget(self, content: RichContent, channel_id: str) -> None:
        """Render HTML to PNG via HtmlRendererPort and upload as inline image."""
        if not self._html_renderer:
            logger.warning("RichContentService: widget received but HtmlRendererPort not configured")
            return

        html = content.data.get("html", "").strip()
        alt_text = content.data.get("alt_text", "Visual card")

        if not html:
            logger.warning("RichContentService: widget missing 'html' field — skipping")
            return

        try:
            png_bytes = await self._html_renderer.render(html)
            await self._media_port.upload_image(
                image_bytes=png_bytes,
                alt_text=alt_text,
                channel_id=channel_id,
            )
        except HtmlRenderError as e:
            logger.error("RichContentService: widget render failed — %s", e)
        except Exception as e:
            logger.error("RichContentService: widget upload failed — %s", e)

    async def _store_html(self, html_content: str, filename: str) -> Optional[str]:
        """Upload HTML to GCS and return public URL. Falls back to plain text upload on error."""
        try:
            html_bytes = html_content.encode("utf-8")
            key = f"html/{uuid.uuid4()}-{filename}"
            url = await self._storage_port.store(
                data=html_bytes,
                key=key,
                content_type="text/html; charset=utf-8",
            )
            return url
        except Exception as e:
            logger.error(
                "RichContentService: GCS upload failed for '%s' — %s; skipping delivery",
                filename,
                e,
            )
            return None


# ------------------------------------------------------------------
# Format converters
# ------------------------------------------------------------------

def _csv_to_xlsx(csv_content: str) -> bytes:
    """Convert CSV string to xlsx bytes via openpyxl."""
    from openpyxl import Workbook  # lazy import — optional dependency

    reader = csv.reader(io.StringIO(csv_content))
    wb = Workbook()
    ws = wb.active
    for row in reader:
        ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _markdown_to_docx(md_content: str) -> bytes:
    """Convert Markdown string to docx bytes via python-docx."""
    import re
    from docx import Document  # lazy import — optional dependency

    doc = Document()
    lines = md_content.splitlines()

    for line in lines:
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _apply_inline(p, re.sub(r"^\d+\. ", "", line).strip())
        elif line.strip() == "---":
            doc.add_paragraph("─" * 40)
        elif line.strip():
            p = doc.add_paragraph()
            _apply_inline(p, line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _apply_inline(paragraph, text: str) -> None:
    """Apply bold/italic inline markdown to a docx paragraph run."""
    import re

    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
